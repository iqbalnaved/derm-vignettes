# === Load DOCX file ===
file_path = "/mnt/d/Naved/Data/derm_vignette/Supplementary Text 3_11032025_highlighted_text.docx"
output_file = "/mnt/d/Naved/Data/derm_vignette/Supplementary_Text_3_11032025_highlighted_text.json"

import re
import json
import os
from docx import Document

# === Load DOCX ===
doc = Document(file_path)

vignettes = []

# === Helper to get highlighted text from a paragraph ===
def extract_highlighted_runs(paragraph):
    highlights = []
    for run in paragraph.runs:
        if run.font.highlight_color:
            text = run.text.strip()
            if text:
                highlights.append(text)
    return highlights

# === Clean vignette text ===
def clean_vignette(text):
    text = re.split(r"['‘’\"]?\s*What is the most likely diagnosis", text, flags=re.IGNORECASE)[0]
    text = re.sub(r"\s+", " ", text).strip()
    return text

# === Process tables ===
for t_i, table in enumerate(doc.tables):
    # Header row contains the diagnosis category
    header_text = table.cell(0, 0).text.strip()
    diag_match = re.search(r'["“]?\s*([A-Za-z\s]+?)\s*["”]?$', header_text)
    diagnosis = diag_match.group(1).strip() if diag_match else "Unknown"

    # Each subsequent row holds case number and vignette
    for r_i, row in enumerate(table.rows[1:], start=1):
        if len(row.cells) < 2:
            continue

        case_num = row.cells[0].text.strip()
        vignette_cell = row.cells[1]

        if not case_num:
            continue

        # Extract highlighted portions from vignette
        highlights = []
        for para in vignette_cell.paragraphs:
            highlights.extend(extract_highlighted_runs(para))

        # Join all highlighted parts into one string (if multiple)
        highlighted_text = " ".join(highlights).strip()
        highlighted_text = re.sub(r"\s+", " ", highlighted_text)

        if highlighted_text:
            vignettes.append({
                "case_number": int(case_num),
                "diagnosis": diagnosis,
                "vignette": highlighted_text
            })

# === Save JSON ===
with open(output_file, "w", encoding="utf-8") as f:
    json.dump(vignettes, f, indent=2, ensure_ascii=False)

print(f"✅ Extracted {len(vignettes)} highlighted-only vignettes and saved to '{output_file}'")

# Preview first few
# for v in vignettes[:3]:
    # print(f"\nCase #{v['case_number']} ({v['diagnosis']}):\n{v['vignette']}\n")

